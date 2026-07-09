from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import db

COLOR_PRIMARY = RGBColor(0x13, 0x2A, 0x4C)
COLOR_ACCENT = RGBColor(0xFF, 0x5A, 0x36)


def _obtener_indicadores():
    sql = """
        SELECT
            (SELECT COUNT(*) FROM Categoria) AS categorias,
            (SELECT COUNT(*) FROM Marca) AS marcas,
            (SELECT COUNT(*) FROM Producto) AS productos,
            (SELECT COUNT(*) FROM Variante_Producto) AS variantes,
            (SELECT COUNT(*) FROM Almacen) AS almacenes,
            (SELECT ISNULL(SUM(cantidad_stock), 0) FROM Inventario) AS unidades_stock,
            (SELECT COUNT(*) FROM Cliente) AS clientes,
            (SELECT COUNT(*) FROM Pedido) AS pedidos,
            (SELECT ISNULL(SUM(total), 0) FROM Pedido) AS ventas_totales,
            (SELECT COUNT(*) FROM Proveedor) AS proveedores,
            (SELECT COUNT(*) FROM Orden_Compra) AS ordenes_compra
    """
    fila = db.query(sql)[0]

    valor_inventario = db.query(
        """
        SELECT ISNULL(SUM(i.cantidad_stock * p.precio_base), 0) AS valor
        FROM Inventario i
        JOIN Variante_Producto vp ON i.id_variante = vp.id_variante
        JOIN Producto p ON vp.id_producto = p.id_producto
        """
    )[0]["valor"]

    cogs = db.query(
        """
        SELECT ISNULL(SUM(dc.cantidad * dc.costo_unitario), 0) AS cogs
        FROM Detalle_Compra dc
        """
    )[0]["cogs"]

    fila["valor_inventario"] = valor_inventario
    fila["cogs"] = cogs
    fila["margen_bruto"] = float(fila["ventas_totales"]) - float(cogs)
    return fila


def _top_productos(limite=5):
    sql = f"""
        SELECT TOP {limite} p.nombre AS producto,
               SUM(dp.cantidad) AS unidades,
               SUM(dp.cantidad * dp.precio_unitario) AS total
        FROM Detalle_Pedido dp
        JOIN Variante_Producto vp ON dp.id_variante = vp.id_variante
        JOIN Producto p ON vp.id_producto = p.id_producto
        GROUP BY p.nombre
        ORDER BY total DESC
    """
    return db.query(sql)


def _stock_critico(limite=10):
    sql = f"""
        SELECT TOP {limite} a.nombre AS almacen, p.nombre AS producto, vp.sku, i.cantidad_stock AS stock
        FROM Inventario i
        JOIN Almacen a ON i.id_almacen = a.id_almacen
        JOIN Variante_Producto vp ON i.id_variante = vp.id_variante
        JOIN Producto p ON vp.id_producto = p.id_producto
        WHERE i.cantidad_stock < ?
        ORDER BY i.cantidad_stock ASC
    """
    return db.query(sql, (db.UMBRAL_STOCK_CRITICO,))


def _estilo_tabla(tabla):
    tabla.style = "Light Grid Accent 1"
    for celda in tabla.rows[0].cells:
        celda.paragraphs[0].runs[0].font.bold = True


def generar_reporte_ejecutivo(ruta_destino):
    ind = _obtener_indicadores()

    doc = Document()

    titulo = doc.add_heading("Reporte Ejecutivo - Total Sport", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_PRIMARY

    doc.add_heading("Indicadores generales", level=1)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = "Indicador"
    tabla.rows[0].cells[1].text = "Valor"
    filas = [
        ("Categorías", ind["categorias"]),
        ("Marcas", ind["marcas"]),
        ("Productos (modelos)", ind["productos"]),
        ("Variantes físicas (SKU)", ind["variantes"]),
        ("Almacenes", ind["almacenes"]),
        ("Unidades en stock", ind["unidades_stock"]),
        ("Valor del inventario (a precio de venta)", f"S/ {float(ind['valor_inventario']):,.2f}"),
        ("Clientes registrados", ind["clientes"]),
        ("Pedidos registrados", ind["pedidos"]),
        ("Ventas acumuladas", f"S/ {float(ind['ventas_totales']):,.2f}"),
        ("Costo de ventas (COGS)", f"S/ {float(ind['cogs']):,.2f}"),
        ("Margen bruto", f"S/ {float(ind['margen_bruto']):,.2f}"),
        ("Proveedores", ind["proveedores"]),
        ("Órdenes de compra", ind["ordenes_compra"]),
    ]
    for etiqueta, valor in filas:
        fila = tabla.add_row()
        fila.cells[0].text = str(etiqueta)
        fila.cells[1].text = str(valor)
    _estilo_tabla(tabla)

    doc.add_heading("Top 5 productos más vendidos", level=1)
    top = _top_productos()
    if top:
        tabla2 = doc.add_table(rows=1, cols=3)
        tabla2.rows[0].cells[0].text = "Producto"
        tabla2.rows[0].cells[1].text = "Unidades vendidas"
        tabla2.rows[0].cells[2].text = "Total vendido (S/)"
        for fila in top:
            r = tabla2.add_row()
            r.cells[0].text = fila["producto"]
            r.cells[1].text = str(fila["unidades"])
            r.cells[2].text = f"{float(fila['total']):,.2f}"
        _estilo_tabla(tabla2)
    else:
        doc.add_paragraph("No hay ventas registradas todavía.")

    doc.add_heading(f"Alertas de stock crítico (menor a {db.UMBRAL_STOCK_CRITICO} unidades)", level=1)
    criticos = _stock_critico()
    if criticos:
        tabla3 = doc.add_table(rows=1, cols=4)
        tabla3.rows[0].cells[0].text = "Almacén"
        tabla3.rows[0].cells[1].text = "Producto"
        tabla3.rows[0].cells[2].text = "SKU"
        tabla3.rows[0].cells[3].text = "Stock"
        for fila in criticos:
            r = tabla3.add_row()
            r.cells[0].text = fila["almacen"]
            r.cells[1].text = fila["producto"]
            r.cells[2].text = fila["sku"]
            r.cells[3].text = str(fila["stock"])
        _estilo_tabla(tabla3)
    else:
        doc.add_paragraph("No hay productos con stock crítico en este momento.")

    doc.save(ruta_destino)
    return ruta_destino
